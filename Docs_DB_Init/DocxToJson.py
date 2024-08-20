import json
import os
from docx import Document
from datetime import datetime

def extract_all_text(doc):
    """Extracts all text from the document, including text in tables."""
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return "\n".join(full_text)

def find_section_content(text, section_name, section_keywords):
    """Finds the content of a section by its name."""
    start_idx = text.find(section_name)
    if start_idx == -1:
        return None
    
    start_idx += len(section_name)
    end_idx = len(text)
    
    for keyword in section_keywords:
        next_section_idx = text.find(keyword, start_idx)
        if next_section_idx != -1 and next_section_idx < end_idx:
            end_idx = next_section_idx

    section_content = text[start_idx:end_idx].strip()
    return section_content if section_content else None

def read_docx(file_path, section_keywords):
    """Reads the .docx file and extracts required sections."""
    doc = Document(file_path)
    full_text = extract_all_text(doc)
    
    sections = {}
    for section_name in section_keywords:
        sections[section_name] = find_section_content(full_text, section_name, section_keywords)
    
    return sections

def get_last_updated(file_path):
    """Gets the last updated date of the document."""
    last_modified_timestamp = os.path.getmtime(file_path)
    last_modified_date = datetime.fromtimestamp(last_modified_timestamp).strftime('%Y-%m-%d')
    return last_modified_date

def convert_to_json(sections, title, last_updated):
    """Converts the extracted sections and additional metadata into JSON format."""
    # Split the title into name and code
    try:
        materia_name, codigo = title.rsplit('_', 1)
    except ValueError:
        materia_name, codigo = title, "Unknown"

    json_data = {
        "version": "1.0",
        "last_updated": last_updated,
        "title": title,
        "description": f"Syllabus de la materia '{materia_name}' con código '{codigo}'",
        "access_control": "privado",
        "content": sections
    }
    return json.dumps(json_data, indent=4, ensure_ascii=False)

def save_json(json_data, output_file):
    """Saves the JSON data to a file."""
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(json_data)

def process_folder(input_folder, output_folder, section_keywords):
    """Processes all .docx files in the input folder and generates JSON files in the output folder."""
    # Obtener la ruta absoluta al script
    script_path = os.path.abspath(__file__)
    # Obtener el directorio del script
    script_dir = os.path.dirname(script_path)
    
    # Construir las rutas absolutas a las carpetas
    input_folder_path = os.path.join(script_dir, input_folder)
    output_folder_path = os.path.join(script_dir, output_folder)

    # Verificar si la carpeta de salida existe, si no, crearla
    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)
    
    # Verificar si la carpeta de entrada existe
    if not os.path.isdir(input_folder_path):
        print(f"La carpeta '{input_folder_path}' no existe.")
        return

    print(f"Verificando archivos en: {input_folder_path}")

    for file_name in os.listdir(input_folder_path):
        if file_name.endswith(".docx"):
            docx_file = os.path.join(input_folder_path, file_name)
            title = os.path.splitext(file_name)[0]
            
            # Read document content
            sections = read_docx(docx_file, section_keywords)
            
            # Get last updated date
            last_updated = get_last_updated(docx_file)
            
            # Convert to JSON
            json_data = convert_to_json(sections, title, last_updated)
            
            # Save JSON to the output folder with the same name as the .docx file
            output_json_file = os.path.join(output_folder_path, f"{title}.json")
            save_json(json_data, output_json_file)
            print(f"Converted {file_name} to {output_json_file}")

if __name__ == "__main__":
    input_folder = "DOCX"  # Path to the folder containing .docx files
    output_folder = "JSON/Syllabus"  # Path to the folder where JSON files will be saved

    # List of section keywords to look for
    section_keywords = [
        "Nombre Corto de la Asignatura",
        "Nombre Largo de la Asignatura",
        "Código de la asignatura",
        "Grado",
        "Descripción",
        "Número de Créditos",
        "Condiciones Académicas de Inscripción (Pre-requisitos)",
        "Período Académico de Vigencia",
        "Objetivos de Formación",
        "Resultados de Aprendizaje Esperados (RAE)",
        "Contenidos temáticos",
        "Estrategias Pedagógicas",
        "Evaluación",
        "Recursos Bibliográficos"
    ]

    # Process all .docx files in the input folder
    process_folder(input_folder, output_folder, section_keywords)
